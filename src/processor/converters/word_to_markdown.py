import re

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
import re
from converters.semantic_split import DocumentChunker
from utils.logging import AppLogger
from converters.title_generate import MarkdownSummarizer  # 新增导入

# 常量定义
logger = AppLogger.get_logger(__name__)
HEADING_REGEX = re.compile(r"^(标题|Heading)\s*([1-6])$")
SEMANTIC_SPLIT_THRESHOLD = 50  # 语义切分阈值
DEFAULT_MAX_HEADING_CHUNK_SIZE = 1000  # 默认标题块最大长度
DEFAULT_FALLBACK_CHUNK_SIZE = 500  # 默认回退块大小
DEFAULT_CHUNK_OVERLAP = 50  # 默认块重叠大小


def iter_block_items(parent):
    for child in parent._element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


# 如果是语义切分, 那么就接受传入的DocumentChunker对象作为语义切分器
def word_to_markdown(word_path, semantic_spliter: DocumentChunker = None):
    doc = Document(word_path)
    md_content = []
    # 跟踪当前标题层级 [h1, h2, h3, h4, h5, h6]
    current_headings = [None, None, None, None, None, None]

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            heading_level = get_heading_level(block)
            if heading_level and 1 <= heading_level <= 6:
                # 更新标题层级
                for i in range(heading_level - 1, 6):
                    current_headings[i] = None
                current_headings[heading_level - 1] = block.text
                
                # 组合标题
                parts = [p for p in current_headings if p is not None]
                combined_text = " | ".join(parts)
                md_content.append("\n----------\n")
                md_content.append(f"### {combined_text}")
            else:
                text = block.text.strip()
                if text:
                    if semantic_spliter and len(text) > SEMANTIC_SPLIT_THRESHOLD:
                        chunks = semantic_spliter.process_text(text)
                        for index, chunk in enumerate(chunks):
                            md_content.append(chunk)
                            if index < len(chunks) - 1:
                                md_content.append("\n----------\n")
                            md_content.append(f"#### 段落块{index}")
                    else:
                        chunks = fixed_length_split(text, DEFAULT_FALLBACK_CHUNK_SIZE, DEFAULT_CHUNK_OVERLAP)
                        for index, chunk in enumerate(chunks):
                            md_content.append(chunk + "  ")  
                            if index < len(chunks) - 1:
                                md_content.append("\n----------\n")
                else:
                    md_content.append("") 

        elif isinstance(block, Table):
            # 处理表格
            rows = []
            for row in block.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")

            if rows:
                # 添加表头分隔线
                col_count = len(block.rows[0].cells)
                separator = "| " + " | ".join(["---"] * col_count) + " |"
                rows.insert(1, separator)
                md_content.append("\n".join(rows))
    return  clean_empty_markdown_chunks("\n".join(md_content))[0]
    # with open(md_path, 'w', encoding='utf-8') as f:
    #     f.write('\n'.join(md_content))


def hybrid_split_to_markdown(
    word_path, max_heading_chunk_size=1000, fallback_chunk_size=500, generate_missing_titles=False
):
    """
    返回 Markdown 格式的切分结果：
    1. 优先按标题切分
    2. 如果某部分过长, 再按固定长度切分
    
    参数:
        generate_missing_titles: 是否为缺少标题的分块生成标题, 默认为True
    """
    # 最后要返回的块集合
    chunks = []
    # 每行md内容
    current_md_lines = []
    # 添加计数器
    heading_only_count = 0
    content_only_count = 0
    generated_title_count = 0  # 生成标题计数器
    # 条件初始化标题生成器
    summarizer = None
    if generate_missing_titles:
        summarizer = MarkdownSummarizer(model_path="./models/models--IDEA-CCNL--Randeng-T5-784M-MultiTask-Chinese", device="cpu")
    # 跟踪当前标题层级 [h1, h2, h3, h4, h5, h6]
    current_headings = [None, None, None, None, None, None]
    # 把文档转换成document对象进行处理
    doc = Document(word_path)
    # 对每个段落进行处理
    for paragraph in doc.paragraphs:
        heading_level = get_heading_level(paragraph)
        if heading_level is not None and 1 <= heading_level <= 6:
            # 更新标题层级
            # 重置当前级别及以下的标题
            for i in range(heading_level - 1, 6):
                current_headings[i] = None
            current_headings[heading_level - 1] = paragraph.text
            # 组合标题
            parts = [p for p in current_headings if p is not None]
            combined_text = " | ".join(parts)
            md_line = f"### {combined_text}"
        else:
            md_line = paragraph_to_markdown(paragraph)
        
        # 如果这个段落元素是docx中的标题
        if is_heading(paragraph):
            if current_md_lines:
                current_chunk = "\n".join(current_md_lines)
                if len(current_chunk) > max_heading_chunk_size:
                    # 过长时按固定长度切分
                    split_chunks = fixed_length_split_md(current_md_lines, fallback_chunk_size)
                    # 修改分块过滤逻辑
                    for chunk in split_chunks:
                        if not has_only_headings(chunk) and not has_only_content(chunk):
                            chunks.append(chunk)
                        else:
                            if has_only_headings(chunk):
                                heading_only_count += 1
                            elif has_only_content(chunk):
                                if generate_missing_titles and summarizer:
                                    # 为缺少标题的分块生成标题
                                    title = summarizer.generate_summary(chunk)
                                    chunk_with_title = f"### {title}\n\n{chunk}"
                                    chunks.append(chunk_with_title)
                                    generated_title_count += 1
                                else:
                                    content_only_count += 1
                else:
                    # 添加过滤逻辑
                    if not has_only_headings(current_chunk) and not has_only_content(current_chunk):
                        chunks.append(current_chunk)
                    else:
                        if has_only_headings(current_chunk):
                            heading_only_count += 1
                        elif has_only_content(current_chunk):
                            content_only_count += 1
                current_md_lines = []
            current_md_lines.append(md_line)
        else:
            current_md_lines.append(md_line)

    # 处理最后一个 chunk
    if current_md_lines:
        final_chunk = "\n".join(current_md_lines)
        if len(final_chunk) > max_heading_chunk_size:
            split_chunks = fixed_length_split_md(current_md_lines, fallback_chunk_size)
            # 添加过滤逻辑
            for chunk in split_chunks:
                if not has_only_headings(chunk) and not has_only_content(chunk):
                    chunks.append(chunk)
                else:
                    if has_only_headings(chunk):
                        heading_only_count += 1
                    elif has_only_content(chunk):
                        if generate_missing_titles and summarizer:
                            # 为缺少标题的分块生成标题
                            title = summarizer.generate_summary(chunk)
                            chunk_with_title = f"### {title}\n\n{chunk}"
                            chunks.append(chunk_with_title)
                            generated_title_count += 1
                        else:
                            content_only_count += 1
        else:
            # 添加过滤逻辑
            if not has_only_headings(final_chunk) and not has_only_content(final_chunk):
                chunks.append(final_chunk)
            else:
                if has_only_headings(final_chunk):
                    heading_only_count += 1
                elif has_only_content(final_chunk):
                    if generate_missing_titles and summarizer:
                        # 为缺少标题的分块生成标题
                        title = summarizer.generate_summary(final_chunk)
                        chunk_with_title = f"### {title}\n\n{final_chunk}"
                        chunks.append(chunk_with_title)
                        generated_title_count += 1
                    else:
                        content_only_count += 1

    # 输出日志
    if heading_only_count > 0:
        logger.info(f"删除了{heading_only_count}个缺少正文的分块")
    if generate_missing_titles:
        if generated_title_count > 0:
            logger.info(f"为{generated_title_count}个缺少标题的分块生成了标题")
    else:
        if content_only_count > 0:
            logger.info(f"删除了{content_only_count}个缺少标题的分块")

    return "\n\n----------\n\n".join(chunks)


def is_heading(paragraph, level=None):
    """检查段落是否是标题及指定级别"""
    if not paragraph.style.name.startswith(("Heading", "标题")):
        return False
    if level is not None:
        # 提取标题级别, 如 "Heading 1" -> 1
        match = re.search(r"\d+", paragraph.style.name)
        if match:
            return int(match.group()) == level
    return True


def fixed_length_split(text, chunk_size=500, overlap=50):
    """按固定长度切分文本, 允许重叠"""
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)
        start += chunk_size - overlap  # 滑动窗口, 允许重叠
    return chunks


def paragraph_to_markdown(paragraph):
    """将单个段落转换为 Markdown 格式"""
    if is_heading(paragraph, level=1):
        return f"# {paragraph.text}"
    elif is_heading(paragraph, level=2):
        return f"## {paragraph.text}"
    elif is_heading(paragraph, level=3):
        return f"### {paragraph.text}"
    elif is_heading(paragraph, level=4):
        return f"#### {paragraph.text}"
    elif is_heading(paragraph, level=5):
        return f"##### {paragraph.text}"
    elif is_heading(paragraph, level=6):
        return f"###### {paragraph.text}"
    elif paragraph.style.name.startswith(("List", "列表")):
        return f"* {paragraph.text}"  # 无序列表
    else:
        return paragraph.text + "  "  # 普通段落, 加两个空格保证换行


def fixed_length_split_md(md_lines, chunk_size=500, overlap=50):
    """对 Markdown 文本按固定长度切分(保留换行符)"""
    full_text = "\n".join(md_lines)
    chunks = []
    start = 0
    while start < len(full_text):
        end = start + chunk_size
        chunk = full_text[start:end]
        chunks.append(chunk)
        start += chunk_size - overlap
    return chunks


def has_only_headings(md_content):
    """检查Markdown内容是否只有标题而没有正文"""
    lines = [line.strip() for line in md_content.split('\n') if line.strip()]
    if not lines:
        return False
    # 检查所有非空行是否都是标题
    return all(line.startswith('#') for line in lines)


def has_only_content(md_content):
    """检查Markdown内容是否只有正文而没有标题"""
    lines = [line.strip() for line in md_content.split('\n') if line.strip()]
    if not lines:
        return False
    # 检查是否存在至少一个标题行
    has_heading = any(line.startswith('#') for line in lines)
    return not has_heading


def get_heading_level(paragraph):
    """获取标题级别(1-6), 非标题返回None

    Args:
        paragraph: docx段落对象

    Returns:
        int: 标题级别(1-6)或None
    """
    match = HEADING_REGEX.match(paragraph.style.name)
    if match:
        return int(match.group(2))
    return None


def clean_empty_markdown_chunks(markdown_content:str=None) -> tuple:
    """
    清理Markdown内容中只有标题没有正文的分块
    分块由若干短划线(-)分隔，删除条件：
    1. 分块内容为空（仅包含空格、换行、制表符等）
    2. 分块中所有非空行都是标题行（以#开头）
    
    参数:
        markdown_content: str, 可选 - 直接提供的Markdown字符串内容
    
    返回:
        tuple: (处理后的Markdown内容, 删除的分块数量)
    """

    content = markdown_content
    
    # 使用正则表达式分割分块（若干短划线分隔）
    chunks = re.split(r'\n\s*----------+\s*\n', content.strip())
    
    cleaned_chunks = []
    deleted_count = 0
    
    for chunk in chunks:
        # 处理分块内容，过滤空白行
        lines = [line.strip() for line in chunk.split('\n') if line.strip()]
        
        # 检查删除条件
        if len(lines) == 0 or all(line.startswith('#') for line in lines):
            deleted_count += 1
            continue
        
        cleaned_chunks.append(chunk)
    
    # 重组分块内容
    cleaned_content = '\n----------\n'.join(cleaned_chunks)
    
    # print(f"清理完成！共删除 {deleted_count} 个缺少正文的Markdown分块")
    return cleaned_content, deleted_count